# docx_exporter.py
print("DOCX_EXPORTER VERSION = NEW_PATCH_20260128_STABLE_TABLE")
import os
import traceback

from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .loader import ProblemItem
from .table_renderer import normalize_table_to_grid, try_find_table
from .config import AppConfig


# ------------------------------------------------------------
# 기본 paragraph 생성 + 문서 스타일 유틸
# ------------------------------------------------------------
def _set_eastasia_font(style, font_name: str):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

def _setup_document(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    _set_eastasia_font(style, "한컴바탕")
    style.font.size = Pt(9)

def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")

def _style_cell_text(cell, size_pt=9, bold=False):
    # ✅ 셀의 문단/런 스타일을 안정적으로 통일
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs:
            r.bold = bold
            r.font.size = Pt(size_pt)

def _add_label(container, text: str):
    p = container.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def _add_problem_header(container, num: int, it: ProblemItem):
    p = container.add_paragraph()
    r1 = p.add_run(f"{num}. ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(f"ID: {it.pid}  ({it.prefix})")
    r2.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def _add_par(doc_or_cell, text: str, bold: bool = False):
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


# ------------------------------------------------------------
# 안정적인 표 생성 (⭐ 씹안정 패치 적용)
# ------------------------------------------------------------
def _add_grid_table(doc_or_cell, table_obj: Any, total_width_in: float = 2.0):
    """
    ✅ 안정 포인트
    - 열 너비는 cell.width가 아니라 t.columns[j].width로만 고정 (Word가 가장 안정적으로 처리)
    - tblCellMar(셀 마진 OXML 조작) 적용하지 않음 (표 씹창의 주범)
    - 텍스트 크기/줄간격만 통일해서 '작고 단정'하게
    """

    headers, rows = normalize_table_to_grid(table_obj)

    if not headers or len(headers) == 0:
        _add_par(doc_or_cell, "(표 데이터 없음 / 파싱 불가)")
        return

    n_cols = len(headers)

    if rows is None or not isinstance(rows, list):
        rows = []

    fixed_rows = []
    for r in rows:
        if r is None:
            r = []
        rr = list(r)
        if len(rr) < n_cols:
            rr = rr + [""] * (n_cols - len(rr))
        elif len(rr) > n_cols:
            rr = rr[:n_cols]
        fixed_rows.append(rr)

    n_rows = 1 + len(fixed_rows)

    t = doc_or_cell.add_table(rows=n_rows, cols=n_cols)
    t.style = "Table Grid"
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    total_width = Inches(total_width_in)
    col_w = int(total_width / n_cols)   # ✅ float → int(EMU)


    # ✅ 열 너비 고정 (중요)
    for j in range(n_cols):
        t.columns[j].width = col_w

    # ---- header ----
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = str(h)
        _style_cell_text(cell, size_pt=9, bold=True)

    # ---- body ----
    for i, row in enumerate(fixed_rows):
        for j in range(n_cols):
            cell = t.rows[i + 1].cells[j]
            v = row[j] if row[j] is not None else ""
            cell.text = str(v)
            _style_cell_text(cell, size_pt=9, bold=False)


# ------------------------------------------------------------
# 텍스트 찾기 함수
# ------------------------------------------------------------
def first_text(payload: Dict[str, Any], keys: Tuple[str, ...] | List[str]) -> Optional[str]:
    for k in keys:
        v = payload.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()
    return None


# ------------------------------------------------------------
# 이미지 처리 (셀/문서 모두 호환)
# ------------------------------------------------------------
def _try_add_image(container, payload: Dict[str, Any]):
    img_path = (
        payload.get("_image_path")
        or payload.get("image_path")
        or payload.get("tree_img")
        or payload.get("figure_path")
        or payload.get("fig_path")
        or payload.get("img")
        or payload.get("diagram_image")
    )

    if not isinstance(img_path, str) or not img_path.strip():
        return

    ip = img_path.strip()

    def resolve_path(p: str) -> Optional[str]:
        if os.path.exists(p):
            return p
        base = os.path.dirname(__file__)
        candidates = [
            os.path.join(base, p),
            os.path.join(base, "..", p),
            os.path.join(base, "..", "..", p),
        ]
        return next((x for x in candidates if os.path.exists(x)), None)

    found = resolve_path(ip)
    if not found:
        _add_par(container, f"(이미지 경로 없음: {ip})")
        return

    try:
        container.add_paragraph("")
        p = container.add_paragraph()
        r = p.add_run()
        r.add_picture(found, width=Inches(1.2))
        container.add_paragraph("")
    except Exception as e:
        _add_par(container, f"(이미지 삽입 실패: {found} / {type(e).__name__})")


# ------------------------------------------------------------
# DOCX EXPORT 메인
# ------------------------------------------------------------
def export_docx_bytes(
    cfg: AppConfig,
    selected: List[ProblemItem],
    include_explanations: bool = True,
    include_full_table: bool = True,
    two_columns: bool = True,
) -> bytes:

    doc = Document()
    _setup_document(doc)

    idx = 0
    pnum = 1

    # --------------------------------------------------------
    # 문제 본문 2단 출력
    # --------------------------------------------------------
    while idx < len(selected):

        if two_columns:
            outer = doc.add_table(rows=1, cols=2)
            outer.autofit = False

            # ✅ outer 테두리만 제거 (레이아웃 안정)
            _remove_table_borders(outer)

            outer.columns[0].width = Inches(3.4)
            outer.columns[1].width = Inches(3.4)

            left = outer.rows[0].cells[0]
            right = outer.rows[0].cells[1]

            targets = [(left, selected[idx], pnum)]
            idx += 1
            pnum += 1
            if idx < len(selected):
                targets.append((right, selected[idx], pnum))
                idx += 1
                pnum += 1
        else:
            targets = [(doc, selected[idx], pnum)]
            idx += 1
            pnum += 1

        for container, it, num in targets:
            payload = it.payload or {}

            # ✅ 헤더를 조금 더 깔끔하게 쓰고 싶으면 아래 1줄로 교체 가능:
            # _add_problem_header(container, num, it)
            _add_par(container, f"[문제 {num}]  ID: {it.pid}  ({it.prefix})", bold=True)

            ptxt = first_text(payload, cfg.problem_text_keys)
            atxt = first_text(payload, cfg.ask_line_keys)

            if ptxt:
                # 더 예쁘게: _add_label(container, "문제")
                _add_par(container, "문제", bold=True)
                _add_par(container, ptxt)

            if atxt:
                # 더 예쁘게: _add_label(container, "요구사항")
                _add_par(container, "요구사항", bold=True)
                _add_par(container, atxt)

            _try_add_image(container, payload)

            given = try_find_table(payload, list(cfg.given_table_keys)) or payload.get("_given_table")
            if given is not None:
                _add_par(container, "제시표", bold=True)
                try:
                    _add_grid_table(container, given, total_width_in=1.8)
                except Exception as e:
                    _add_par(container, f"(표 변환 실패: {type(e).__name__}: {e})")
                    print("TABLE FAIL:", type(e).__name__, e)
                    print(traceback.format_exc())

        if idx < len(selected):
            doc.add_page_break()

    # --------------------------------------------------------
    # 정답/해설 파트
    # --------------------------------------------------------
    doc.add_page_break()
    _add_par(doc, "[정답/해설]", bold=True)

    for i, it in enumerate(selected, start=1):
        payload = it.payload or {}
        _add_par(doc, f"{i}. ID: {it.pid}  ({it.prefix})", bold=True)

        ans = first_text(payload, cfg.answer_keys)
        expl = first_text(payload, cfg.explanation_keys)

        _add_par(doc, "정답", bold=True)
        _add_par(doc, ans or "(없음)")

        if include_full_table:
            full = try_find_table(payload, list(cfg.full_table_keys)) or payload.get("_full_table")
            if full is not None:
                _add_par(doc, "완성표", bold=True)
                try:
                    _add_grid_table(doc, full, total_width_in=6.4)
                except Exception:
                    _add_par(doc, "(완성표 변환 실패)")

        if include_explanations:
            _add_par(doc, "해설", bold=True)
            _add_par(doc, expl or "(없음)")

        _add_par(doc, "-" * 40)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
